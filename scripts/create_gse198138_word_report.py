from __future__ import annotations

import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "outputs" / "GSE198138_processed"
FIG_DIR = BASE / "figures"
DOCX_PATH = BASE / "GSE198138_DrugReflector_中文图文报告.docx"

COLORS = {
    "navy": "#0B1F33",
    "blue": "#2166AC",
    "light_blue": "#9ECAE1",
    "red": "#B2182B",
    "orange": "#D95F02",
    "green": "#1B9E77",
    "gray": "#6B7280",
    "light_gray": "#E5E7EB",
    "purple": "#5E3C99",
}


def setup_matplotlib() -> None:
    plt.rcParams.update(
        {
            "font.family": "Arial",
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "axes.linewidth": 0.8,
            "axes.edgecolor": "#111827",
            "axes.labelcolor": "#111827",
            "xtick.color": "#374151",
            "ytick.color": "#374151",
            "axes.titlesize": 9,
            "axes.labelsize": 8,
            "xtick.labelsize": 7,
            "ytick.labelsize": 7,
            "legend.fontsize": 7,
            "figure.dpi": 160,
            "savefig.dpi": 320,
            "savefig.bbox": "tight",
        }
    )


def read_inputs() -> tuple[pd.DataFrame, dict[str, pd.DataFrame], pd.DataFrame, pd.DataFrame]:
    meta = pd.read_csv(BASE / "expression" / "GSE198138_sample_metadata.csv")
    if "Unnamed: 0" in meta.columns:
        meta = meta.drop(columns=["Unnamed: 0"])
    diff_tables = {}
    for label in [
        "FXS_vs_Unaffected_all",
        "FXS_vs_Unaffected_day22",
        "FXS_vs_Unaffected_day42_48",
    ]:
        diff_tables[label] = pd.read_csv(
            BASE / "differential" / f"GSE198138_{label}_derived_differential_table.csv"
        )
    predictions = pd.read_csv(BASE / "predictions" / "GSE198138_online_api_predictions_top50.csv")
    signature = pd.read_csv(BASE / "differential" / "GSE198138_DrugReflector_signature_matrix.csv", index_col=0)
    return meta, diff_tables, predictions, signature


def save_figure(fig: plt.Figure, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    png = FIG_DIR / f"{name}.png"
    pdf = FIG_DIR / f"{name}.pdf"
    fig.savefig(png, transparent=False)
    fig.savefig(pdf, transparent=False)
    plt.close(fig)
    return png


def plot_pipeline() -> Path:
    fig, ax = plt.subplots(figsize=(7.2, 2.1))
    ax.axis("off")
    steps = [
        ("GEO archive", "analysis.tar.gz\n.cloupe.gz"),
        ("Derived expression", "cluster mean UMI\nweighted by cells"),
        ("Gene mapping", "Ensembl ID\nHGNC symbol"),
        ("Signature", "FXS vs\nUnaffected"),
        ("DrugReflector", "compound\nprioritization"),
    ]
    x_positions = np.linspace(0.07, 0.93, len(steps))
    for i, (title, body) in enumerate(steps):
        x = x_positions[i]
        box = plt.Rectangle(
            (x - 0.085, 0.35),
            0.17,
            0.38,
            facecolor="#F8FAFC",
            edgecolor=COLORS["navy"],
            linewidth=0.9,
            transform=ax.transAxes,
        )
        ax.add_patch(box)
        ax.text(x, 0.63, title, ha="center", va="center", fontsize=8, fontweight="bold", color=COLORS["navy"], transform=ax.transAxes)
        ax.text(x, 0.47, body, ha="center", va="center", fontsize=7, color="#374151", transform=ax.transAxes)
        if i < len(steps) - 1:
            ax.annotate(
                "",
                xy=(x_positions[i + 1] - 0.105, 0.54),
                xytext=(x + 0.105, 0.54),
                arrowprops=dict(arrowstyle="->", lw=1.0, color=COLORS["gray"]),
                xycoords=ax.transAxes,
                textcoords=ax.transAxes,
            )
    ax.text(
        0.5,
        0.16,
        "Exploratory workflow based on public Cell Ranger summary files; not a raw single-cell count reanalysis.",
        ha="center",
        va="center",
        fontsize=7.5,
        color=COLORS["red"],
        transform=ax.transAxes,
    )
    return save_figure(fig, "Figure_1_workflow")


def plot_sample_design(meta: pd.DataFrame) -> Path:
    counts = (
        meta.groupby(["differentiation_day", "disease_status"])
        .size()
        .unstack(fill_value=0)
        .reindex(["day_22", "day_42-48"])
    )
    cells = (
        meta.groupby(["differentiation_day", "disease_status"])["n_cells_graphclust"]
        .sum()
        .unstack(fill_value=0)
        .reindex(["day_22", "day_42-48"])
    )
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 2.8))
    for ax, frame, title, ylabel in [
        (axes[0], counts, "Sample design", "Samples"),
        (axes[1], cells, "Cells represented in graphclust", "Cells"),
    ]:
        bottom = np.zeros(len(frame))
        for status, color in [("Unaffected", COLORS["blue"]), ("FXS", COLORS["red"])]:
            values = frame.get(status, pd.Series(0, index=frame.index)).to_numpy()
            ax.bar(frame.index, values, bottom=bottom, label=status, color=color, width=0.55)
            bottom += values
        ax.set_title(title, loc="left", fontweight="bold")
        ax.set_ylabel(ylabel)
        ax.spines[["top", "right"]].set_visible(False)
        ax.legend(frameon=False)
    return save_figure(fig, "Figure_2_sample_design")


def plot_mapping_summary(signature: pd.DataFrame) -> Path:
    input_genes = 22926
    mapped = signature.shape[1]
    unmapped = input_genes - mapped
    fig, ax = plt.subplots(figsize=(4.2, 3.0))
    ax.bar(["Mapped", "Unmapped"], [mapped, unmapped], color=[COLORS["green"], COLORS["light_gray"]], width=0.58)
    ax.set_ylabel("Genes")
    ax.set_title("Gene mapping coverage", loc="left", fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    rate = mapped / input_genes * 100
    ax.text(0, mapped * 1.02, f"{mapped:,}\n({rate:.1f}%)", ha="center", va="bottom", fontsize=8, color=COLORS["green"])
    ax.text(1, max(unmapped, 1) * 1.05, f"{unmapped:,}", ha="center", va="bottom", fontsize=8, color=COLORS["gray"])
    return save_figure(fig, "Figure_3_gene_mapping")


def plot_signature_heatmap(signature: pd.DataFrame) -> Path:
    top_genes = signature.abs().max(axis=0).sort_values(ascending=False).head(35).index
    mat = signature[top_genes]
    fig, ax = plt.subplots(figsize=(7.2, 4.5))
    vmax = np.nanpercentile(np.abs(mat.to_numpy()), 97)
    im = ax.imshow(mat.to_numpy(), aspect="auto", cmap="RdBu_r", vmin=-vmax, vmax=vmax)
    ax.set_yticks(np.arange(mat.shape[0]))
    ax.set_yticklabels([idx.replace("FXS_vs_Unaffected_", "") for idx in mat.index])
    ax.set_xticks(np.arange(mat.shape[1]))
    ax.set_xticklabels(mat.columns, rotation=90)
    ax.set_title("High-amplitude signature genes", loc="left", fontweight="bold")
    cbar = fig.colorbar(im, ax=ax, fraction=0.025, pad=0.02)
    cbar.set_label("v-score")
    ax.tick_params(length=0)
    return save_figure(fig, "Figure_4_signature_heatmap")


def plot_volcano_panels(diff_tables: dict[str, pd.DataFrame]) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(9.6, 3.1), sharey=True)
    for ax, (label, table) in zip(axes, diff_tables.items()):
        work = table.copy()
        p = pd.to_numeric(work["pvalue_welch_on_derived_sample_means"], errors="coerce").clip(lower=1e-300)
        y = -np.log10(p)
        x = pd.to_numeric(work["log2FC_FXS_vs_Unaffected"], errors="coerce")
        score = pd.to_numeric(work["v_score_FXS_vs_Unaffected"], errors="coerce")
        colors = np.where(score > 0, COLORS["red"], COLORS["blue"])
        ax.scatter(x, y, s=5, c=colors, alpha=0.35, linewidths=0)
        ax.axvline(0, color="#111827", lw=0.6)
        ax.set_title(label.replace("FXS_vs_Unaffected_", ""), loc="left", fontweight="bold")
        ax.set_xlabel("log2FC")
        ax.spines[["top", "right"]].set_visible(False)
        top = work.reindex(score.abs().sort_values(ascending=False).head(4).index)
        for _, row in top.iterrows():
            ax.text(
                row["log2FC_FXS_vs_Unaffected"],
                -math.log10(max(row["pvalue_welch_on_derived_sample_means"], 1e-300)) + 0.05,
                str(row["gene"]),
                fontsize=5.8,
                color=COLORS["navy"],
            )
    axes[0].set_ylabel("-log10(P value)")
    return save_figure(fig, "Figure_5_effect_size_panels")


def plot_top_gene_bars(diff_tables: dict[str, pd.DataFrame]) -> Path:
    table = diff_tables["FXS_vs_Unaffected_all"].copy()
    up = table.sort_values("v_score_FXS_vs_Unaffected", ascending=False).head(12)
    down = table.sort_values("v_score_FXS_vs_Unaffected", ascending=True).head(12)
    plot_df = pd.concat([down, up], axis=0)
    fig, ax = plt.subplots(figsize=(5.6, 5.2))
    colors = [COLORS["blue"] if value < 0 else COLORS["red"] for value in plot_df["v_score_FXS_vs_Unaffected"]]
    ax.barh(plot_df["gene"], plot_df["v_score_FXS_vs_Unaffected"], color=colors)
    ax.axvline(0, color="#111827", lw=0.7)
    ax.set_xlabel("v-score")
    ax.set_title("Top FXS-associated signature genes", loc="left", fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    return save_figure(fig, "Figure_6_top_signature_genes")


def _display_label(row: pd.Series) -> str:
    for key in ["display_name", "compound"]:
        value = row.get(key)
        if pd.notna(value) and str(value).strip():
            return str(value)
    return "Unknown"


def plot_compound_bars(predictions: pd.DataFrame) -> list[Path]:
    paths = []
    for label, table in predictions.groupby("signature"):
        top = table.sort_values("rank").head(15).copy()
        top["label"] = top.apply(_display_label, axis=1)
        top = top.iloc[::-1]
        fig, ax = plt.subplots(figsize=(5.6, 4.8))
        gradient = plt.cm.Blues(np.linspace(0.38, 0.9, len(top)))
        ax.barh(top["label"], top["probability"], color=gradient)
        ax.set_xlabel("Predicted probability")
        ax.set_title(label.replace("FXS_vs_Unaffected_", "Top compounds: "), loc="left", fontweight="bold")
        ax.spines[["top", "right"]].set_visible(False)
        paths.append(save_figure(fig, f"Figure_7_compounds_{label}"))
    return paths


def plot_moa_summary(predictions: pd.DataFrame) -> Path:
    top = predictions[predictions["rank"] <= 50].copy()
    top = top[top["moa"].notna() & top["moa"].astype(str).str.strip().ne("")]
    exploded = top.assign(moa=top["moa"].astype(str).str.split(r"\s*\|\s*", regex=True)).explode("moa")
    counts = exploded["moa"].value_counts().head(12).iloc[::-1]
    fig, ax = plt.subplots(figsize=(6.0, 4.2))
    ax.barh(counts.index, counts.values, color=COLORS["purple"])
    ax.set_xlabel("Occurrences in top-50 lists")
    ax.set_title("Mechanism-of-action recurrence", loc="left", fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    return save_figure(fig, "Figure_8_moa_summary")


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = "" if pd.isna(text) else str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)
            run.bold = bold
            set_run_font(run, "Microsoft YaHei")


def set_run_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        set_run_font(run)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(75, 85, 99)
    set_run_font(run)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int = 15) -> None:
    table_df = df[columns].head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for cell, column in zip(cells, columns):
            value = row[column]
            if isinstance(value, float):
                value = f"{value:.4g}"
            set_cell_text(cell, value)


def safe_fill(value: object) -> str:
    if pd.isna(value) or str(value).strip() == "":
        return "未匹配到客观注释"
    return str(value)


def create_docx(
    meta: pd.DataFrame,
    diff_tables: dict[str, pd.DataFrame],
    predictions: pd.DataFrame,
    signature: pd.DataFrame,
    figures: dict[str, Path | list[Path]],
) -> None:
    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    title = doc.add_heading("GSE198138 DrugReflector 中文图文分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于公开 GEO 单细胞汇总文件的探索性转录组 signature 药物优先级筛选")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)
    set_run_font(run)

    doc.add_heading("一、结论摘要", level=1)
    add_bullets(
        doc,
        [
            "本报告分析对象为 GSE198138，该数据集为人源 hPSC 分化 GABAergic inhibitory neurons 的 Fragile X Syndrome（FXS）与 Unaffected 对照单细胞转录组研究。",
            "GEO 公开文件未提供标准 cell-by-gene count matrix，因此本报告没有伪装成完整单细胞重分析；本次结果来自 Cell Ranger graphclust cluster mean UMI 与 cluster cell counts 的样本级派生表达。",
            f"本次共形成 14 个样本、{signature.shape[1]:,} 个 HGNC 基因的 derived expression / signature 输入，并构建 3 个对比：all、day22、day42-48。",
            "DrugReflector 输出代表基于转录组 signature 的候选化合物优先级，不等同于直接靶向结合证据，也不等同于临床疗效证据。",
        ],
    )

    doc.add_heading("二、数据来源与可用性审核", level=1)
    add_paragraph(
        doc,
        "GSE198138 的公开补充文件包括 GSE198138_RAW.tar，其中每个样本主要包含 Cell Ranger analysis.tar.gz 和 Loupe Browser 使用的 .cloupe.gz 文件。"
        "本地检查未发现 matrix.mtx、features.tsv、barcodes.tsv、filtered_feature_bc_matrix.h5 或 .h5ad，因此不能客观还原标准单细胞 count matrix。"
    )
    add_paragraph(
        doc,
        "因此，本报告采用保守策略：只使用 Cell Ranger 明确导出的 graphclust clusters.csv 和 differential_expression.csv 中的 cluster mean UMI，"
        "按每个 cluster 的细胞数加权，得到每个样本的 gene-level derived average expression，再计算 FXS 相对 Unaffected 的 v-score / log2FC。"
    )
    doc.add_picture(str(figures["workflow"]), width=Cm(16.2))
    add_caption(doc, "图 1. 数据处理流程。图中所有节点均来自公开可追溯文件；本流程属于探索性派生分析。")

    doc.add_heading("三、样本设计与基因映射", level=1)
    doc.add_picture(str(figures["sample_design"]), width=Cm(16.2))
    add_caption(doc, "图 2. 样本与 graphclust 细胞量分布。FXS 与 Unaffected 同时覆盖 day 22 与 day 42-48。")
    doc.add_picture(str(figures["mapping"]), width=Cm(11.0))
    add_caption(doc, "图 3. Ensembl gene ID 到 HGNC symbol 的映射覆盖度。")
    add_table(
        doc,
        meta,
        ["geo_accession", "disease_status", "differentiation_day", "n_cells_graphclust"],
        ["GEO样本", "分组", "分化时间", "graphclust细胞数"],
        max_rows=20,
    )

    doc.add_heading("四、差异 signature 结果", level=1)
    add_paragraph(
        doc,
        "方向定义：v-score 或 log2FC 为正表示 FXS 高于 Unaffected；负值表示 FXS 低于 Unaffected。"
        "由于本分析基于派生样本均值，p 值与 BH 校正仅作为探索性参考，不应作为正式统计显著性结论。"
    )
    doc.add_picture(str(figures["heatmap"]), width=Cm(16.2))
    add_caption(doc, "图 4. 三个对比中变化幅度较高的 signature genes。")
    doc.add_picture(str(figures["volcano"]), width=Cm(16.2))
    add_caption(doc, "图 5. 三个对比的 log2FC 与名义 P 值分布；标注基于绝对 v-score 较高的基因。")
    doc.add_picture(str(figures["top_genes"]), width=Cm(13.5))
    add_caption(doc, "图 6. 全样本 FXS vs Unaffected 的 Top up/down signature genes。")

    for label, table in diff_tables.items():
        doc.add_heading(f"{label} 候选基因", level=2)
        add_paragraph(doc, "Top 15 FXS-up genes：")
        add_table(
            doc,
            table.sort_values("v_score_FXS_vs_Unaffected", ascending=False),
            ["gene", "v_score_FXS_vs_Unaffected", "log2FC_FXS_vs_Unaffected", "padj_BH_exploratory"],
            ["基因", "v-score", "log2FC", "探索性BH校正"],
        )
        add_paragraph(doc, "Top 15 FXS-down genes：")
        add_table(
            doc,
            table.sort_values("v_score_FXS_vs_Unaffected", ascending=True),
            ["gene", "v_score_FXS_vs_Unaffected", "log2FC_FXS_vs_Unaffected", "padj_BH_exploratory"],
            ["基因", "v-score", "log2FC", "探索性BH校正"],
        )

    doc.add_heading("五、DrugReflector 候选化合物筛选", level=1)
    add_paragraph(
        doc,
        "DrugReflector 结果表示输入 signature 与训练集中化合物扰动转录组模式之间的模型匹配优先级。"
        "当前部署未配置 CLUE Touchstone signed connectivity API，因此 direction 字段统一显示 No objective evidence；这意味着不能客观判定候选化合物是逆转还是模拟输入 signature。"
    )
    for path in figures["compound_bars"]:
        doc.add_picture(str(path), width=Cm(13.8))
        add_caption(doc, "图 7. 对应 signature 的 Top 15 candidate compounds。条形长度为模型预测概率。")
    doc.add_picture(str(figures["moa"]), width=Cm(14.0))
    add_caption(doc, "图 8. Top-50 候选化合物中重复出现的 mechanism-of-action 类别。")

    for label, table in predictions.groupby("signature", sort=False):
        doc.add_heading(f"{label} Top 候选化合物", level=2)
        work = table.sort_values("rank").head(15).copy()
        for col in ["display_name", "target", "moa", "chemical_name"]:
            work[col] = work[col].map(safe_fill)
        add_table(
            doc,
            work,
            ["rank", "compound", "display_name", "probability", "target", "moa"],
            ["排名", "Broad编号", "名称", "预测概率", "客观靶点注释", "客观MOA注释"],
            max_rows=15,
        )

    doc.add_heading("六、机制解释与研究建议", level=1)
    add_paragraph(
        doc,
        "从全样本结果看，retinoid receptor agonist 相关候选物（如 AM-580、TTNPB、tamibarotene、tazarotene）在全样本与晚期时间点结果中反复出现，"
        "提示 retinoid/RAR 轴可能与该 FXS-derived signature 存在较强的转录组模式关联。该结论仍是 signature-level 关联，不能直接解释为 RARA/RARB/RARG 已被证明是疾病核心靶点。"
    )
    add_paragraph(
        doc,
        "day22 结果与 day42-48 结果差异较明显，说明 FXS 转录组扰动可能存在发育阶段依赖性。若用于后续实验设计，建议分时间点验证，而不是把所有时间点合并后直接给出单一机制结论。"
    )
    add_bullets(
        doc,
        [
            "优先补齐标准原始数据：从 SRA FASTQ 重新跑 Cell Ranger / STARsolo，获得真正的 cell-by-gene count matrix。",
            "在标准 count matrix 基础上重新做 QC、doublet 过滤、批次校正、细胞类型注释、pseudo-bulk 差异分析。",
            "候选药物验证建议先聚焦全样本与 day42-48 中重复出现、且有客观化学名/靶点/MOA 注释的化合物。",
            "若要判断药物是逆转还是模拟 FXS signature，需要接入 CLUE/LINCS signed connectivity 或自行构建化合物扰动表达数据集。",
            "实验验证不应逐个盲试全部 Top 50，而应按机制类别、可获得性、细胞毒性、已有 CNS/神经发育证据进行分层筛选。",
        ],
    )

    doc.add_heading("七、交付文件", level=1)
    file_rows = pd.DataFrame(
        [
            ["派生表达矩阵", "expression/GSE198138_derived_sample_expression_HGNC.h5ad"],
            ["DrugReflector输入signature", "differential/GSE198138_DrugReflector_signature_matrix.csv"],
            ["差异表", "differential/GSE198138_*_derived_differential_table.csv"],
            ["线上API预测结果", "predictions/GSE198138_online_api_predictions_top50.csv"],
            ["本报告配图", "figures/*.png 和 figures/*.pdf"],
        ],
        columns=["内容", "相对路径"],
    )
    add_table(doc, file_rows, ["内容", "相对路径"], ["内容", "相对路径"], max_rows=10)

    core_properties = doc.core_properties
    core_properties.title = "GSE198138 DrugReflector 中文图文分析报告"
    core_properties.subject = "Exploratory signature-based compound prioritization"
    core_properties.author = "DrugReflector analysis pipeline"
    doc.save(DOCX_PATH)


def main() -> None:
    setup_matplotlib()
    meta, diff_tables, predictions, signature = read_inputs()
    figures = {
        "workflow": plot_pipeline(),
        "sample_design": plot_sample_design(meta),
        "mapping": plot_mapping_summary(signature),
        "heatmap": plot_signature_heatmap(signature),
        "volcano": plot_volcano_panels(diff_tables),
        "top_genes": plot_top_gene_bars(diff_tables),
        "compound_bars": plot_compound_bars(predictions),
        "moa": plot_moa_summary(predictions),
    }
    create_docx(meta, diff_tables, predictions, signature, figures)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
