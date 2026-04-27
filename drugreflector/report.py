"""Word report generation for DrugReflector results."""

from __future__ import annotations

import io
import math
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.request import urlopen

import matplotlib

matplotlib.use("Agg")

import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _configure_matplotlib_font() -> fm.FontProperties | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for path in candidates:
        if not path.exists():
            continue
        prop = fm.FontProperties(fname=str(path))
        matplotlib.rcParams["font.family"] = "sans-serif"
        matplotlib.rcParams["font.sans-serif"] = [prop.get_name(), "DejaVu Sans", "Arial Unicode MS"]
        matplotlib.rcParams["axes.unicode_minus"] = False
        return prop
    matplotlib.rcParams["axes.unicode_minus"] = False
    return None


MATPLOTLIB_FONT = _configure_matplotlib_font()


@dataclass
class ReportContext:
    sample: str
    locale: str
    meta: dict[str, Any]
    rows: list[dict[str, Any]]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _apply_axis_font(ax) -> None:
    if MATPLOTLIB_FONT is None:
        return
    for title in (ax.title, getattr(ax, "_left_title", None), getattr(ax, "_right_title", None)):
        if title is not None:
            title.set_fontproperties(MATPLOTLIB_FONT)
    x_label = ax.xaxis.label
    if x_label is not None:
        x_label.set_fontproperties(MATPLOTLIB_FONT)
    y_label = ax.yaxis.label
    if y_label is not None:
        y_label.set_fontproperties(MATPLOTLIB_FONT)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(MATPLOTLIB_FONT)


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None:
            return default
        if isinstance(value, str) and not value.strip():
            return default
        return float(value)
    except Exception:
        return default


def _get_display_name(row: dict[str, Any]) -> str:
    annotation = row.get("annotation") or {}
    return str(annotation.get("display_name") or row.get("compound") or "Unknown")


def _get_direction_label(direction: dict[str, Any] | None, locale: str) -> str:
    label = str((direction or {}).get("label") or "").strip().lower()
    if label == "reverse":
        return "更像逆转" if locale == "zh" else "Reverse"
    if label == "mimic":
        return "更像增强" if locale == "zh" else "Mimic"
    return "暂无客观证据" if locale == "zh" else "No objective evidence"


def _get_direction_summary(direction: dict[str, Any] | None, locale: str) -> str:
    label = str((direction or {}).get("label") or "").strip().lower()
    if label == "reverse":
        return (
            "该化合物更像逆转输入 signature，但这仍是签名层面的数据证据，不等于已经证实某个单基因被上调或下调。"
            if locale == "zh"
            else "This compound appears to reverse the input signature, but that remains signature-level evidence rather than proof of up- or down-regulation of any single gene."
        )
    if label == "mimic":
        return (
            "该化合物更像增强或模拟输入 signature，但这仍是签名层面的数据证据，不等于已经证实某个单基因被上调或下调。"
            if locale == "zh"
            else "This compound appears to mimic the input signature, but that remains signature-level evidence rather than proof of up- or down-regulation of any single gene."
        )
    return (
        "当前部署缺少可客观判定方向性的外部签名连通性证据，因此不能负责任地判断为“更像逆转”或“更像增强”。"
        if locale == "zh"
        else "This deployment lacks objective external signed-connectivity evidence, so it cannot responsibly classify the compound as reverse or mimic."
    )


def _format_ratio(value: Any) -> str:
    numeric = _safe_float(value, float("nan"))
    if math.isnan(numeric):
        return "不可用"
    return f"{numeric * 100:.1f}%"


def _summarize_tokens(rows: list[dict[str, Any]], key: str, limit: int = 8) -> list[tuple[str, int]]:
    counts: dict[str, int] = {}
    for row in rows:
        annotation = row.get("annotation") or {}
        raw = annotation.get(key)
        if not raw:
            continue
        for token in str(raw).split("|"):
            clean = token.strip()
            if not clean:
                continue
            counts[clean] = counts.get(clean, 0) + 1
    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))[:limit]


def _annotation_tier(row: dict[str, Any]) -> str:
    annotation = row.get("annotation") or {}
    score = 0
    if annotation.get("display_name"):
        score += 1
    if annotation.get("target") or annotation.get("moa"):
        score += 1
    if annotation.get("pubchem_cid") or annotation.get("structure_image"):
        score += 1
    if score >= 3:
        return "High"
    if score == 2:
        return "Medium"
    return "Low"


def _top_rows(rows: list[dict[str, Any]], n: int = 25) -> list[dict[str, Any]]:
    return sorted(rows, key=lambda row: _safe_float(row.get("prob")), reverse=True)[:n]


def _render_bar_chart(rows: list[dict[str, Any]], locale: str) -> io.BytesIO:
    top = list(reversed(_top_rows(rows, 15)))
    labels = [_get_display_name(row)[:28] for row in top]
    values = [_safe_float(row.get("prob")) for row in top]
    colors = ["#1f6fd5" if idx >= len(top) - 3 else "#78b8f4" for idx in range(len(top))]
    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    ax.barh(labels, values, color=colors, edgecolor="none")
    ax.set_xlabel("预测概率" if locale == "zh" else "Predicted probability")
    ax.set_title("Top compounds" if locale != "zh" else "候选化合物排名图", loc="left", fontsize=13, fontweight="bold")
    ax.grid(axis="x", linestyle="--", alpha=0.25)
    ax.set_axisbelow(True)
    _apply_axis_font(ax)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=220, facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _render_scatter(rows: list[dict[str, Any]], locale: str) -> io.BytesIO:
    top = _top_rows(rows, 50)
    ranks = [_safe_float(row.get("rank")) for row in top]
    probs = [_safe_float(row.get("prob")) for row in top]
    logits = [_safe_float(row.get("logit")) for row in top]
    labels = [_get_display_name(row) for row in top]
    fig, ax = plt.subplots(figsize=(7.8, 4.8))
    sizes = [50 + max(logit, 0.0) * 8 for logit in logits]
    colors = ["#1f6fd5" if rank <= 10 else "#9cccf7" for rank in ranks]
    ax.scatter(ranks, probs, s=sizes, c=colors, alpha=0.85, edgecolors="white", linewidths=0.6)
    for rank, prob, label in zip(ranks[:5], probs[:5], labels[:5]):
        ax.text(
            rank + 0.35,
            prob,
            label[:20],
            fontsize=8,
            fontproperties=MATPLOTLIB_FONT,
        )
    ax.set_xlabel("排名" if locale == "zh" else "Rank")
    ax.set_ylabel("预测概率" if locale == "zh" else "Predicted probability")
    ax.set_title("全局视图" if locale == "zh" else "Global scatter view", loc="left", fontsize=13, fontweight="bold")
    ax.grid(True, linestyle="--", alpha=0.25)
    _apply_axis_font(ax)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=220, facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _render_coverage(rows: list[dict[str, Any]], locale: str) -> io.BytesIO:
    top = _top_rows(rows, 20)
    levels = ["High", "Medium", "Low"]
    counts = [sum(1 for row in top if _annotation_tier(row) == level) for level in levels]
    fig, ax = plt.subplots(figsize=(5.8, 4.0))
    ax.bar(levels, counts, color=["#2e8b57", "#e6a93c", "#c9d2dc"], edgecolor="none")
    ax.set_title("注释覆盖度" if locale == "zh" else "Annotation coverage", loc="left", fontsize=13, fontweight="bold")
    ax.set_ylabel("数量" if locale == "zh" else "Count")
    for idx, count in enumerate(counts):
        ax.text(
            idx,
            count + 0.15,
            str(count),
            ha="center",
            fontsize=9,
            fontproperties=MATPLOTLIB_FONT,
        )
    _apply_axis_font(ax)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=220, facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _render_token_chart(title: str, items: list[tuple[str, int]]) -> io.BytesIO | None:
    if not items:
        return None
    labels = [item[0][:30] for item in reversed(items)]
    counts = [item[1] for item in reversed(items)]
    fig, ax = plt.subplots(figsize=(7.6, max(2.4, 0.45 * len(items) + 1.2)))
    ax.barh(labels, counts, color="#5ca8e8", edgecolor="none")
    ax.set_title(title, loc="left", fontsize=13, fontweight="bold")
    ax.grid(axis="x", linestyle="--", alpha=0.25)
    ax.set_axisbelow(True)
    _apply_axis_font(ax)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=220, facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _fetch_structure_image(url: str | None) -> io.BytesIO | None:
    if not url:
        return None
    try:
        with urlopen(url, timeout=20) as resp:
            data = resp.read()
        buf = io.BytesIO(data)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_picture_with_caption(document: Document, image: io.BytesIO | None, caption: str) -> None:
    if image is None:
        return
    document.add_picture(image, width=Inches(6.4))
    p = document.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)


def _build_overview_table(document: Document, ctx: ReportContext) -> None:
    meta = ctx.meta
    input_quality = meta.get("input_quality") or {}
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("输入来源", str(meta.get("filename") or ctx.sample)),
        ("样本 / signature", ctx.sample),
        ("候选化合物数", str(meta.get("n_compounds", "不可用"))),
        ("输入基因数", str(meta.get("n_genes", "不可用"))),
        ("模型基因重叠率", _format_ratio(input_quality.get("overlap_ratio"))),
        ("模型基因重叠数", f"{input_quality.get('overlap_gene_count', '不可用')} / {input_quality.get('model_gene_count', '不可用')}"),
    ]
    ortholog = meta.get("ortholog_mapping")
    if ortholog:
        rows.extend(
            [
                ("跨物种映射来源", str(ortholog.get("source") or "不可用")),
                (
                    "小鼠→人映射成功率",
                    _format_ratio(
                        _safe_float(ortholog.get("mapped_input_genes")) / max(_safe_float(ortholog.get("input_genes")), 1.0)
                    ),
                ),
            ]
        )
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        _set_cell_shading(cells[0], "EAF2FB")


def _build_top_compounds_table(document: Document, rows: list[dict[str, Any]], locale: str) -> None:
    table = document.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["排名", "化合物", "预测概率", "Logit", "方向性", "靶点", "机制"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header if locale == "zh" else ["Rank", "Compound", "Probability", "Logit", "Direction", "Target", "MOA"][idx]
        _set_cell_shading(cell, "DCEBFA")
    for row in _top_rows(rows, 15):
        annotation = row.get("annotation") or {}
        cells = table.add_row().cells
        cells[0].text = str(row.get("rank", ""))
        cells[1].text = _get_display_name(row)
        cells[2].text = f"{_safe_float(row.get('prob')):.4f}"
        cells[3].text = f"{_safe_float(row.get('logit')):.4f}"
        cells[4].text = _get_direction_label(row.get("direction"), locale)
        cells[5].text = str(annotation.get("target") or "暂无客观注释" if locale == "zh" else "Unavailable")
        cells[6].text = str(annotation.get("moa") or "暂无客观注释" if locale == "zh" else "Unavailable")


def _add_selected_compound_section(document: Document, row: dict[str, Any], locale: str) -> None:
    annotation = row.get("annotation") or {}
    direction = row.get("direction") or {}
    name = _get_display_name(row)
    _add_heading(document, f"重点化合物解析：{name}", level=2)
    _add_paragraph(document, f"{name} 为当前结果中优先级最高的候选化合物，原始编号为 {row.get('compound', 'Unknown')}。")
    bullets = [
        f"排名：#{row.get('rank', '不可用')}；预测概率：{_safe_float(row.get('prob')):.4f}；Logit：{_safe_float(row.get('logit')):.4f}。",
        f"方向性：{_get_direction_label(direction, locale)}。{_get_direction_summary(direction, locale)}",
        f"客观注释名称：{annotation.get('display_name') or '暂无客观注释'}；化学名称：{annotation.get('chemical_name') or '暂无客观注释'}。",
        f"靶点：{annotation.get('target') or '暂无客观注释'}；作用机制：{annotation.get('moa') or '暂无客观注释'}。",
        f"临床阶段：{annotation.get('clinical_phase') or '暂无客观注释'}；适应症领域：{annotation.get('disease_area') or '暂无客观注释'}。",
    ]
    for bullet in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(bullet)
    image = _fetch_structure_image(annotation.get("structure_image"))
    if image is not None:
        _add_picture_with_caption(
            document,
            image,
            "图 5. 重点化合物的二维化学结构图，来自客观可核对的 PubChem 结构记录。",
        )


def _add_limitations(document: Document, ctx: ReportContext) -> None:
    _add_heading(document, "结果解读边界与注意事项", level=2)
    notes = [
        "本报告属于基于转录组 signature 的候选药物优先级筛选，不等同于直接蛋白结合证据，也不等同于临床有效性结论。",
        "方向性仅指化合物相对于输入 signature 更像逆转、增强，或暂无客观证据；这不直接等于某个目标基因已经被上调或下调。",
        "如果本次输入涉及小鼠到人的同源映射，应优先关注最终进入模型的基因覆盖率，而不是单纯用原始跨物种映射率判断结果优劣。",
        "建议把排名靠前且机制收敛的化合物作为后续文献核对、细胞实验、qPCR、Western blot 或功能实验的优先验证对象。",
    ]
    for note in notes:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(note)


def build_docx_report(response: dict[str, Any], sample: str | None = None, locale: str = "zh") -> bytes:
    data = response.get("data") or {}
    meta = response.get("meta") or {}
    samples = data.get("samples") or []
    if not samples:
        raise ValueError("Prediction response does not contain any sample results.")
    selected_sample = sample or str(samples[0])
    results = data.get("results") or {}
    rows = list(results.get(selected_sample) or [])
    if not rows:
        raise ValueError(f"Prediction response does not contain results for sample '{selected_sample}'.")

    ctx = ReportContext(sample=selected_sample, locale=locale, meta=meta, rows=rows)
    document = Document()
    _set_document_style(document)

    title = "DrugReflector 中文分析报告" if locale == "zh" else "DrugReflector Analysis Report"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(str(meta.get("filename") or selected_sample))

    _add_heading(document, "一、结果总览", level=1)
    _add_paragraph(
        document,
        "本报告基于当前页面结果自动生成，所有名称、靶点、机制、结构和方向性说明仅在存在客观数据时显示；无可靠来源处均保留为“暂无客观注释”或“暂无客观证据”。",
    )
    _build_overview_table(document, ctx)

    _add_heading(document, "二、图形化结果", level=1)
    _add_picture_with_caption(
        document,
        _render_bar_chart(rows, locale),
        "图 1. 候选化合物排名图。条形长度代表模型给出的预测概率，越靠前表示优先级越高。",
    )
    _add_paragraph(
        document,
        "解析：该图用于快速识别优先级最高的候选化合物。需要注意的是，高排名代表模型优先级更高，并不直接等于该化合物已被实验确认有效。",
    )
    _add_picture_with_caption(
        document,
        _render_scatter(rows, locale),
        "图 2. 全局散点图。横轴为排名，纵轴为预测概率，点的大小与 Logit 强弱相关。",
    )
    _add_paragraph(
        document,
        "解析：该图用于判断候选化合物是否形成明显的头部梯队。如果前列化合物和后续化合物之间存在较大间距，通常说明模型对头部候选的区分度更强。",
    )
    _add_picture_with_caption(
        document,
        _render_coverage(rows, locale),
        "图 3. 注释覆盖度。展示前 20 个候选化合物中，高、中、低注释完整度的分布情况。",
    )
    _add_paragraph(
        document,
        "解析：注释覆盖度越高，说明该候选化合物越容易进入后续机制分析和文献验证。仅有 BRD 编号而无客观注释的条目仍可保留，但解释空间会更有限。",
    )

    moa_chart = _render_token_chart("图 4A. 机制聚合" if locale == "zh" else "Figure 4A. MOA aggregation", _summarize_tokens(_top_rows(rows, 20), "moa"))
    if moa_chart is not None:
        _add_picture_with_caption(
            document,
            moa_chart,
            "图 4A. 前 20 个候选化合物中客观记录到的作用机制聚合情况。",
        )
    target_chart = _render_token_chart("图 4B. 靶点聚合" if locale == "zh" else "Figure 4B. Target aggregation", _summarize_tokens(_top_rows(rows, 20), "target"))
    if target_chart is not None:
        _add_picture_with_caption(
            document,
            target_chart,
            "图 4B. 前 20 个候选化合物中客观记录到的靶点聚合情况。",
        )
    _add_paragraph(
        document,
        "解析：如果多个头部候选收敛到相似的靶点或机制，这通常比单独盯住某一个药物更有解释价值，因为它提示了潜在的稳定通路信号。",
    )

    _add_heading(document, "三、Top 候选化合物表", level=1)
    _build_top_compounds_table(document, rows, locale)
    _add_paragraph(
        document,
        "解析：表格把排名、概率、方向性、靶点和机制放在一起，方便直接作为内部汇报或科研记录的核心结果页。",
    )

    _add_heading(document, "四、重点化合物深入解析", level=1)
    _add_selected_compound_section(document, _top_rows(rows, 1)[0], locale)

    _add_heading(document, "五、综合结论", level=1)
    top_names = "、".join(_get_display_name(row) for row in _top_rows(rows, 5))
    _add_paragraph(
        document,
        f"综合当前模型输出，优先级靠前的候选化合物主要包括：{top_names}。这些条目应被视为后续验证的优先名单，而不是已经成立的治疗结论。",
    )
    _add_paragraph(
        document,
        "建议优先结合头部候选的客观靶点和作用机制，判断它们是否在生物学上收敛到同一通路；如存在机制收敛，可优先设计更聚焦的验证实验。",
    )

    _add_limitations(document, ctx)

    document.add_section(WD_SECTION.NEW_PAGE)
    _add_heading(document, "附录：原始导出字段说明", level=1)
    appendix = [
        "Probability：模型对当前候选化合物给出的相对优先级分数。",
        "Logit：模型原始打分，通常与概率排序一致，但不是生物实验中的效应值。",
        "Direction：仅表示化合物相对输入 signature 的方向关系，不直接等于单个目标基因的上调或下调。",
        "Target / MOA / Structure：仅在 Broad Repurposing Hub 与 PubChem 有客观匹配记录时显示。",
    ]
    for item in appendix:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()
